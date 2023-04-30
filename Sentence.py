import tkinter
import tkinter.messagebox
import shutil
from tkinter import messagebox
import customtkinter
import docx
import zipfile, os
from tkinter import filedialog

customtkinter.set_appearance_mode("System") 
customtkinter.set_default_color_theme("green") 

def unzip_docx(docx_path, output_path):
    with zipfile.ZipFile(docx_path, 'r') as zip_file:
        zip_file.extractall(output_path)
def zip_directory(folder_path, output_path):
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zip_file.write(file_path, os.path.relpath(file_path, folder_path))

    shutil.rmtree(folder_path, ignore_errors=False, onerror=None)
class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        self.title("Sentence")
        self.geometry(f"{1100}x{580}")
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure((2, 3), weight=0)
        self.grid_rowconfigure((0, 1, 2), weight=1)
        self.logo_label = customtkinter.CTkLabel(self, text="Sentence", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(2, 0))
        self.sidebar_button_2 = customtkinter.CTkButton(self, command=self.open, text="Open File")
        self.sidebar_button_2.grid(row=3, column=0, padx=20, pady=10)
        self.sidebar_button_2 = customtkinter.CTkButton(self, command=self.savelol, text="Save File")
        self.sidebar_button_2.grid(row=4, column=0, padx=20, pady=10)
        self.entry = customtkinter.CTkEntry(self, placeholder_text="Title")
        self.entry.grid(row=3, column=1, columnspan=2, padx=(20, 0), pady=(20, 20), sticky="nsew")
        self.textbox = customtkinter.CTkTextbox(self, width=250, height=500)
        self.textbox.grid(row=0, column=1, padx=(20, 0), pady=(20, 0), sticky="nsew")
        self.textbox.insert("0.0", "Your text")



    def savelol(self):
        result = messagebox.askyesno("Sentence", "Do you want to save with Sentece Openfile feature? (May cause lil problems with Word, Is encrypted (useless))")
        file_patha = filedialog.asksaveasfilename(filetypes=[('Word Documents', '*.docx')])
        filename = os.path.basename(file_patha)
        print(file_patha)
        open(file_patha, 'w')
        doc = docx.Document()
        noracist = self.entry.get()
        racist = self.textbox.get(1.0, "end")
        doc.add_heading(noracist, 0)
        doc.add_paragraph(racist)
        if not file_patha.endswith(".docx"):
            nothing = file_patha
            file_path = nothing + ".docx"
        else:
            file_path = file_patha
        doc.save(file_path)
        if result:
            temp_dir_path = os.path.join(os.path.dirname(file_path), "Sentence_Temp")
            unzip_docx(file_path, temp_dir_path)
            os.remove(file_path)
            with open(os.path.join(temp_dir_path, "sentence.txt"), "w") as f:
                f.write(f"TITLE: {noracist}\nTEXT: {racist}")
            zip_directory(temp_dir_path, file_path)
    def open(self):
        file_pathe = filedialog.askopenfile(filetypes=[('Word Documents', '*.docx')])
        file_path = file_pathe.name
        file_patho = file_pathe.name
        # check if the file path ends with the desired extension
        if file_path.endswith(".docx"):
            print()
        else:
            file_path += ".docx"
            pass
        with zipfile.ZipFile(file_patho, 'r') as zip_ref:
            if 'sentence.txt' in zip_ref.namelist():
                with zip_ref.open('sentence.txt') as file:
                    # Read the lines of the file as bytes objects
                    lines = file.readlines()

                    # Decode the bytes objects into strings using the appropriate encoding
                    title = lines[0].decode('utf-8').split(': ')[1].strip()
                    text = lines[1].decode('utf-8').split(': ')[1].strip()

                    self.textbox.delete(1.0, "end")
                    self.textbox.insert("end", text)
                    self.entry.delete(0, "end")
                    self.entry.insert(0, title)

            else:
                messagebox.showerror("Sentence", "Sorry! This file wasnt saved with the openfile option or isnt from Sentence!")



if __name__ == "__main__":
    print("Welcome to sentence!\nLightweight tool\nSaves to .docx")
    app = App()
    app.mainloop()
